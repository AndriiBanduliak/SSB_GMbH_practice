import io
import re
from datetime import date
import logging

import PyPDF2 # Для работы с PDF-файлами
from docx import Document # Для создания и чтения DOCX-файлов
from docx.enum.text import WD_ALIGN_PARAGRAPH # Для выравнивания текста в DOCX

from config import Translations # Для получения всех переведенных текстов и заголовков
from utils import number_to_ukrainian_words # Для перевода чисел в слова

# Инициализируем логгер для этого модуля
logger = logging.getLogger(__name__)

def extract_text_from_file(file_content: bytes, file_extension: str) -> str:
    """
    Извлекает текстовое содержимое из байтового представления файла.
    Поддерживает форматы TXT, PDF, DOCX.

    :param file_content: Содержимое файла в виде байтовой строки.
    :param file_extension: Расширение файла (например, '.txt', '.pdf', '.docx').
    :return: Извлеченный текст из файла.
    :raises ValueError: Если формат файла не поддерживается или произошла
                        ошибка при извлечении текста (например, файл поврежден).
    """
    try:
        if file_extension == '.txt':
            # Для текстовых файлов просто декодируем байты в строку
            return file_content.decode('utf-8', errors='ignore')
        elif file_extension == '.pdf':
            # Для PDF-файлов используем PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text: # Убеждаемся, что страница содержит текст
                    text += page_text + "\n"
            if not text.strip(): # Если текст не удалось извлечь, возможно, PDF - это картинка
                raise ValueError("No extractable text found in PDF (might be scanned image).")
            return text
        elif file_extension == '.docx':
            # Для DOCX-файлов используем python-docx
            doc = Document(io.BytesIO(file_content))
            return "\n".join(p.text for p in doc.paragraphs)
        else:
            logger.warning(f"Unsupported file format attempted: {file_extension}")
            raise ValueError(Translations.get_text('uk', "error_file_format")) # Use a translation key here

    except Exception as e:
        logger.error(f"Error extracting text from {file_extension} file: {e}", exc_info=True)
        # Более общее сообщение, чтобы не раскрывать внутренние детали ошибки
        raise ValueError(Translations.get_text('uk', "error_file_extract"))


def generate_advocate_request_doc(request_data: dict, lang_code: str, ai_response_text: str) -> io.BytesIO:
    """
    Генерирует документ адвокатского запроса в формате .docx на основе предоставленных данных
    и текста, сгенерированного AI.

    :param request_data: Словарь с данными, собранными в процессе диалога для запроса.
    :param lang_code: Код языка для получения правильных переводов.
    :param ai_response_text: Основной текст запроса, сгенерированный AI.
    :return: Объект io.BytesIO, содержащий сгенерированный .docx файл.
    """
    doc = Document()
    T = lambda key: Translations.get_text(lang_code, key) # Удобная функция для сокращения вызовов Translations

    # Заголовок документа
    doc.add_heading("АДВОКАТСЬКИЙ ЗАПИТ", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")

    # Получатель
    doc.add_paragraph(f"До: {request_data.get('recipient_details', 'N/A')}")
    doc.add_paragraph("\n")

    # Отправитель (Адвокат)
    doc.add_paragraph("Від:")
    doc.add_paragraph(f"- {T('prompt_header_form')}: {request_data.get('legal_form_text', 'N/A')}")
    if request_data.get('bureau_name') and request_data['bureau_name'] != "N/A":
        doc.add_paragraph(f"- {T('prompt_header_bureau')}: {request_data.get('bureau_name')}")
    doc.add_paragraph(f"- {T('prompt_header_advocate_name')}: {request_data.get('advocate_name', 'N/A')}")
    doc.add_paragraph(f"- {T('prompt_header_address')}: {request_data.get('mailing_address', 'N/A')}")
    doc.add_paragraph(f"- {T('prompt_header_phone')}: {request_data.get('advocate_phone', 'N/A')}")
    doc.add_paragraph(f"- {T('prompt_header_email')}: {request_data.get('advocate_email', 'N/A')}")
    doc.add_paragraph(f"- {T('prompt_header_certificate')}: {request_data.get('certificate_details', '')}, {request_data.get('certificate_issuer', '')}")
    doc.add_paragraph(f"- {T('prompt_header_order')}: {request_data.get('order_details', 'N/A')}")
    doc.add_paragraph("\n")

    # Клиент, в чьих интересах действует адвокат
    client_type_header = T('prompt_header_client_type_phys') if request_data.get('client_type') == 'phys' else T('prompt_header_client_type_legal')
    client_id_info = f", Код ЄДРПОУ: {request_data.get('client_id', 'N/A')}" if request_data.get('client_type') == 'legal' else ""
    client_full_info = f"{client_type_header}: {request_data.get('client_name', 'N/A')}{client_id_info}, Адреса: {request_data.get('client_address', 'N/A')}"
    doc.add_paragraph(f"В інтересах: {client_full_info}")
    doc.add_paragraph("\n")

    # Основание для оказания правовой помощи
    doc.add_paragraph("Підстава:")
    doc.add_paragraph(f"- {T('prompt_header_contract')}: {request_data.get('contract_details', 'N/A')}")
    doc.add_paragraph(f"- {T('prompt_header_subject')}: {request_data.get('legal_aid_subject', 'N/A')}")
    doc.add_paragraph("\n")
    
    # Детали запроса
    doc.add_paragraph("Деталі запиту:")
    doc.add_paragraph(f"- {T('prompt_header_number')}: {request_data.get('outgoing_number', 'N/A')}")
    doc.add_paragraph(f"- {T('prompt_header_date')}: {date.today().strftime('%d.%m.%Y')}")
    doc.add_paragraph("\n")

    # Основной текст запроса, сгенерированный AI
    # Перевод AI-текста должен происходить в AI-сервисе или промпте, а не здесь
    # Документы Docx не понимают HTML-теги <b>, <i>, <code>, их нужно убрать или заменить на .bold/.italic
    ai_response_cleaned = ai_response_text.replace('<b>', '').replace('</b>', '').replace('<i>', '').replace('</i>', '').replace('<code>', '').replace('</code>', '')
    doc.add_paragraph(ai_response_cleaned)
    doc.add_paragraph("\n")

    # Пункт об ответственности (статичный текст)
    liability_clause = T(f'advocate_request_liability_clause_{lang_code}')
    doc.add_paragraph(liability_clause)
    doc.add_paragraph("\n")

    # Подпись
    doc.add_paragraph("З повагою,")
    doc.add_paragraph(request_data.get('advocate_name', 'Адвокат')) # ФИО адвоката
    doc.add_paragraph("_________________________ (підпис)") # Место для подписи
    doc.add_paragraph("\n")

    # Переводящий и сохраняющий DOCX в байтовый поток
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0) # Перемещаем курсор в начало потока для чтения
    return file_stream


def generate_contract_doc(contract_data: dict, lang_code: str) -> io.BytesIO:
    """
    Генерирует документ "Договор о предоставлении правовой помощи" в формате .docx.
    Динамически вставляет данные клиента, адвоката, условия оплаты и другие детали.

    :param contract_data: Словарь со всеми собранными данными для договора.
    :param lang_code: Код языка для перевода статичных текстов.
    :return: Объект io.BytesIO, содержащий сгенерированный .docx файл.
    """
    doc = Document()
    T = lambda key: Translations.get_text(lang_code, key) # Удобная функция для сокращения вызовов Translations

    # Заголовок документа
    doc.add_heading(T('contract_header'), level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")

    # Номер и дата договора
    p_num_date = doc.add_paragraph()
    p_num_date.add_run(f"Договір про надання правової допомоги №{contract_data['contract_number']}").bold = True
    
    # Форматирование даты
    date_template = T('contract_date_header_template').format(
        day=contract_data['contract_day'],
        month=contract_data['contract_month'],
        year=contract_data['contract_year']
    )
    p_num_date.add_run(date_template)
    p_num_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("\n") # Отступ

    # 2. Стороны Договора
    doc.add_heading(T('contract_client_details_header').replace('*', '').replace('---', ''), level=2)
    client_name = contract_data['client']['name']
    client_position = contract_data['client']['position']
    client_fio = contract_data['client']['fio']
    client_basis = contract_data['client']['basis']
    doc.add_paragraph(f"**{client_name}**, у подальшому «КЛІЄНТ», в особі **{client_position}** {client_fio}, що діє на підставі {client_basis}, з однієї сторони,")

    doc.add_heading(T('contract_advocate_details_header').replace('*', '').replace('---', ''), level=2)
    advocate_fio = contract_data['advocate']['fio']
    advocate_cert_series = contract_data['advocate']['cert_series']
    advocate_cert_number = contract_data['advocate']['cert_number']
    advocate_cert_issuer = contract_data['advocate']['cert_issuer']
    advocate_decision_date_number = contract_data['advocate']['decision_date_number']

    # Парсим дату и номер решения из объединенной строки
    decision_date_part = "N/A"
    decision_number_part = "N/A"
    if advocate_decision_date_number:
        parts = advocate_decision_date_number.split('№')
        if len(parts) == 2:
            decision_date_part = parts[0].strip()
            decision_number_part = parts[1].strip()
        elif parts: # Если только дата, без номера
            decision_date_part = parts[0].strip()

    cert_info = ""
    if advocate_cert_issuer:
        cert_info += f", виданого на підставі рішення {advocate_cert_issuer}"
        if decision_date_part != "N/A":
            cert_info += f" від {decision_date_part}"
        if decision_number_part != "N/A":
            cert_info += f" №{decision_number_part}"

    doc.add_paragraph(
        f"та **Адвокат {advocate_fio}**, далі «АДВОКАТ», що діє на підставі **Свідоцтва про право на заняття адвокатською діяльністю серії {advocate_cert_series} №{advocate_cert_number}{cert_info}**, з іншої сторони, уклали цей договір про наступне:")
    doc.add_paragraph("\n")

    # 1. Предмет Договора
    doc.add_heading(T('contract_section_1_header'), level=2)
    doc.add_paragraph("1.1. КЛІЄНТ доручає, а АДВОКАТ приймає на себе зобов’язання надавати правову допомогу в обсязі та на умовах, передбачених даним Договором.")
    doc.add_paragraph("1.2. Правова допомога полягає у:")
    list_items_1_2 = [
        "- наданні усних консультацій, порад та роз’яснень щодо правових питань (позицій);",
        "- зборі інформації, документів, матеріалів, та їх правовий аналіз;",
        "- складанні процесуальних документів;",
        "- виконанні окремих доручень КЛІЄНТА, що не зашкодять його інтересам;",
        "- участі у судових засіданнях;",
        "- представництві, захисті прав та інтересів КЛІЄНТА в органах державної влади та місцевого самоврядування, в установах, організаціях і підприємствах усіх форм власності, у судах всіх інстанцій."
    ]
    for item in list_items_1_2:
        doc.add_paragraph(item, style='List Bullet')

    # 2. Обязанности и права сторон
    doc.add_heading(T('contract_section_2_header'), level=2)
    # Тексты берутся напрямую из Translations. Они уже не содержат *маркдаун*,
    # а используют HTML <b>, который здесь надо убрать.
    doc.add_paragraph(T('section_2_content_2_1').replace('<b>', '').replace('</b>', ''))
    doc.add_paragraph(T('section_2_content_2_2').replace('<b>', '').replace('</b>', ''))
    doc.add_paragraph(T('section_2_content_2_3').replace('<b>', '').replace('</b>', ''))
    doc.add_paragraph(T('section_2_content_2_4').replace('<b>', '').replace('</b>', ''))

    # 3. Оплата и порядок расчетов (Динамическая часть)
    doc.add_heading(T('contract_section_3_header'), level=2)
    # Текст об оплате уже сгенерирован и форматирован в handler-ах
    doc.add_paragraph(contract_data['payment_details']['text'].replace('<b>', '').replace('</b>', ''))

    # 4. Конфиденциальность
    doc.add_heading(T('contract_section_4_header'), level=2)
    doc.add_paragraph("4.1. КЛІЄНТ та АДВОКАТ зобов’язуються суворо дотримуватися режиму конфіденційності щодо отриманої один від одного інформації та здійснюватимуть всі можливі заходи для попередження розголошення отриманої інформації.")
    doc.add_paragraph("4.2. Обсяг адвокатської таємниці, що не підлягає розголошенню, встановлюється Законом України ”Про адвокатуру та адвокатську діяльність”, відповідними нормативними актами, даним Договором.")
    doc.add_paragraph("4.3. АДВОКАТ не несе відповідальності за розголошення предмета адвокатської таємниці у випадку, якщо таке розголошення було здійснено у відповідності з чинним законодавством України або за згодою КЛІЄНТА.")

    # 5. Форс-мажор
    doc.add_heading(T('contract_section_5_header'), level=2)
    doc.add_paragraph("5.1. Сторони звільняються (частково чи повністю) від зобов’язань за договором у випадку неможливості їх виконання через пожежу, військові дії, землетруси, розпорядження та рішення Верховної Ради України, Президента України, Кабінету Міністрів України, інших державних органів, та інших незалежних від Сторін обставин.")
    doc.add_paragraph("5.2. Сторона, що посилається на обставини непереборної сили, зобов’язана письмово повідомити про їх настання іншу сторону не пізніше 48 годин.")
    doc.add_paragraph("5.3. АДВОКАТ звільняється від відповідальності за спричинені негативні для КЛІЄНТА правові наслідки, що виникли у зв’язку з набуттям чинності змін в законодавстві, про які АДВОКАТ не міг знати та передбачити.")

    # 6. Срок действия Договора
    doc.add_heading(T('contract_section_6_header'), level=2)
    doc.add_paragraph(f"6.1. Даний Договір укладений на строк до {contract_data['end_date']} та набирає чинності з моменту його підписання.")
    doc.add_paragraph("6.2. Цей Договір може бути достроково припинений за взаємною згодою Сторін або розірваний на вимогу однієї із Сторін на умовах, передбачених Договором, Законом України „Про адвокатуру та адвокатську діяльність”, Правилами адвокатської етики.")
    doc.add_paragraph("При цьому, КЛІЄНТ зобов’язаний оплатити АДВОКАТУ гонорар за всю роботу, що була виконана чи підготовлена до виконання, а АДВОКАТ зобов’язаний повідомити КЛІЄНТА про можливі наслідки та ризики, пов’язані з достроковим припиненням Договору.")

    # 7. Изменение условий
    doc.add_heading(T('contract_section_7_header'), level=2)
    doc.add_paragraph("7.1. Умови даного Договору мають однакову обов’язкову силу для Сторін та можуть бути змінені за взаємною домовленістю з обов’язковим складенням письмового документа.")

    # 8. Прочие условия
    doc.add_heading(T('contract_section_8_header'), level=2)
    doc.add_paragraph("8.1. КЛІЄНТ дає згоду АДВОКАТУ на обробку своїх персональних даних.")
    doc.add_paragraph("8.2. КЛІЄНТ дає згоду на отримання ним звернень (повідомлень, звітів, запитів тощо) засобами СМС–розсилок, поштового зв’язку, електронною поштою, телефонним та/або факсимільним зв’язком.")
    doc.add_paragraph("8.3. Даний Договір складений в двох оригінальних примірниках, по одному для кожної із Сторін, кожен з яких має однакову юридичну силу.")
    doc.add_paragraph("8.4. У випадках, не передбачених даним Договором, Сторони керуються чинним законодавством України.")

    # 9. Реквизиты и подписи сторон
    doc.add_heading(T('contract_section_9_header'), level=2)
    
    # Реквизиты Адвоката
    advocate_heading = doc.add_paragraph()
    advocate_heading.add_run("АДВОКАТ").bold = True
    doc.add_paragraph(advocate_fio)
    doc.add_paragraph(f"Свідоцтво про право на заняття адвокатською діяльністю серії {advocate_cert_series} №{advocate_cert_number}{cert_info}")
    doc.add_paragraph(f"Місцезнаходження: {contract_data['advocate']['location']}")
    doc.add_paragraph("\n") # Отступ
    doc.add_paragraph(f"___________________ {advocate_fio}") # Подпись: ФИО адвоката
    doc.add_paragraph("\n")

    # Реквизиты Клиента
    client_heading = doc.add_paragraph()
    client_heading.add_run("КЛІЄНТ").bold = True
    doc.add_paragraph(client_name)
    doc.add_paragraph(f"{client_position} {client_fio}")
    doc.add_paragraph(f"ЄДРПОУ {contract_data['client']['edrpou']}")
    doc.add_paragraph(f"Місцезнаходження: {contract_data['client']['location']}")
    doc.add_paragraph("\n") # Отступ
    doc.add_paragraph(f"___________________ {client_fio}") # Подпись: ФИО клиента
    
    # Сохраняем документ в байтовый поток
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0) # Перемещаем курсор в начало потока для чтения
    return file_stream