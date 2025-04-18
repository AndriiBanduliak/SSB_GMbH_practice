import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

def add_hyperlink(paragraph, text, url):
    """Adds a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), "0000FF"); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t'); text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# Full resume content in German and English
resumes = {
    "de": {
        "name": "Andrii Banduliak",
        "phone": "+49 160 91441767",
        "email": "aobanduliak@gmail.com",
        "linkedin": "https://linkedin.com/in/andrii-banduliak",
        "github": "https://github.com/AndriiBanduliak",
        "summary_heading": "Profil",
        "summary": (
            "Motivierter Softwareentwickler und Data Engineer mit juristischem Hintergrund, "
            "spezialisiert auf Python, Cloud Computing (GCP) und MLOps. Leidenschaftlich "
            "datengetriebene Lösungen und Automatisierung vorantreibend, untermauert durch "
            "umfangreiche Erfahrung in rechtlicher Recherche und Vertragsmanagement."
        ),
        "sections_order": [
            "Kernkompetenzen", "Berufserfahrung", "Ausgewählte Projekte",
            "Bildung", "Zertifikate", "Sprachen", "Schwerbehinderung", "Ehrenamt & Interessen"
        ],
        "sections": {
            "Kernkompetenzen": [
                "Programmiersprachen: Python, SQL, JavaScript, Java",
                "Cloud & Infrastruktur: Google Cloud Platform (Compute Engine, BigQuery, Cloud Composer), Docker, Kubernetes",
                "Data Engineering & MLOps: Apache Airflow, Kubeflow, ML-Pipelines",
                "Frameworks & Tools: TensorFlow, PyTorch, Django, Flask, Git, GitHub Actions, Jenkins",
                "Datenbanken: PostgreSQL, MySQL, MongoDB",
                "Soft Skills: Analytisches Denken, Projektmanagement, interdisziplinäre Kommunikation"
            ],
            "Berufserfahrung": [
                "Freelance Softwareentwickler & Data Engineer (April 2023 – heute): Entwicklung und Wartung von ML-Pipelines mit Kubeflow auf GCP; Automatisierung von Workflows mittels Airflow und GitHub Actions (-30% manueller Aufwand); Discord-Bot für Echtzeitanalyse; PDF-zu-Audio & Speech-to-Text.",
                "Mitbegründer & Rechtsanwalt – Ledovskoy & Banduliak (Mai 2018 – Mai 2023): Leitung der Kanzlei; Beratung in Zivil-, Handels- und Verwaltungsrecht; Vertragsentwurf, Verhandlung, Compliance."
            ],
            "Ausgewählte Projekte": [
                "SSB_GMbH_practice (Python): Automatisierte Analyseplattform für Geschäftsprozesse.",
                "find_word_by_lvl_in_text (Python): Semantische Textanalyse zur Wortklassifizierung.",
                "helper_to_learn_lang (Python/Django): Web-App für das Erlernen von Sprachen.",
                "mlops-on-gcp (Jupyter Notebooks): Erweiterung und Anpassung von Kubeflow-Pipelines auf GCP."
            ],
            "Bildung": [
                "Master of Laws (LL.M.) – Nationale Akademie für Innere Angelegenheiten, Ukraine (2011–2013)"
            ],
            "Zertifikate": [
                "Google Cloud Digital Leader (2024)",
                "IBM Data Science Professional Certificate (2024)",
                "Meta Back-End Developer Professional Certificate (2024)",
                "Coursera: Data Engineering Foundations, GANs, Cloud Engineer Professional (2024)",
                "edX: CS50, Introduction to Computer Science with Python (2023)"
            ],
            "Sprachen": [
                "Ukrainisch, Russisch (Muttersprache)",
                "Englisch (fortgeschritten)",
                "Deutsch (A2)"
            ],
            "Schwerbehinderung": [
                    "Grad der Behinderung: 80 % G B aG"


            ],
            "Ehrenamt & Interessen": [
                "Freiwillige Unterstützung bei Präsidentschaftswahlen (2014) und EURO 2012.",
                "Engagement in Open-Source-Projekten und Teilnahme an Tech-Meetups."
            ]
        }
    },
    "en": {
        "name": "Andrii Banduliak",
        "phone": "+49 160 91441767",
        "email": "aobanduliak@gmail.com",
        "linkedin": "https://linkedin.com/in/andrii-banduliak",
        "github": "https://github.com/AndriiBanduliak",
        "summary_heading": "Profile",
        "summary": (
            "Motivated Software Developer and Data Engineer with a legal background, "
            "specializing in Python, Google Cloud Platform (GCP), and MLOps. Passionate "
            "about data-driven solutions and automation, backed by extensive experience "
            "in legal research and contract management."
        ),
        "sections_order": [
            "Core Competencies", "Work Experience", "Selected Projects",
            "Education", "Certificates", "Languages", "Disability", "Volunteer & Interests"
        ],
        "sections": {
            "Core Competencies": [
                "Programming Languages: Python, SQL,JavaScript, Java",
                "Cloud & Infrastructure: Google Cloud Platform (Compute Engine, BigQuery, Cloud Composer), Docker, Kubernetes",
                "Data Engineering & MLOps: Apache Airflow, Kubeflow, ML Pipelines",
                "Frameworks & Tools: TensorFlow, PyTorch, Django, Flask, Git, GitHub Actions, Jenkins",
                "Databases: PostgreSQL, MySQL, MongoDB",
                "Soft Skills: Analytical Thinking, Project Management, Cross-disciplinary Communication"
            ],
            "Work Experience": [
                "Freelance Software Developer & Data Engineer (April 2023 – Present): Developed ML pipelines with Kubeflow on GCP; Automated workflows (-30% manual effort); Built Discord bot; PDF-to-Audio & Speech-to-Text.",
                "Co-founder & Attorney – Ledovskoy & Banduliak (May 2018 – May 2023): Led law firm; Advised on civil, commercial, and administrative law; Drafted contracts; Ensured compliance."
            ],
            "Selected Projects": [
                "SSB_GMbH_practice (Python): Automated business process analysis platform.",
                "find_word_by_lvl_in_text (Python): Semantic text analysis for word classification.",
                "helper_to_learn_lang (Python/Django): Web app for language learning assistance.",
                "mlops-on-gcp (Jupyter Notebooks): Adapted Kubeflow pipelines on GCP."
            ],
            "Education": [
                "Master of Laws (LL.M.) – National Academy of Internal Affairs, Ukraine (2011–2013)"
            ],
            "Certificates": [
                "Google Cloud Digital Leader (2024)",
                "IBM Data Science Professional Certificate (2024)",
                "Meta Back-End Developer Professional Certificate (2024)",
                "Coursera: Data Engineering Foundations, GANs, Cloud Engineer Professional (2024)",
                "edX: CS50, Introduction to Computer Science with Python (2023)"
            ],
            "Languages": [
                "Ukrainian, Russian (native)",
                "English (advanced)",
                "German (A2)"
            ],
            "Disability": [
                "Degree of disability: 80% (severe)"
            ],
            "Volunteer & Interests": [
                "Volunteered at presidential elections (2014) and EURO 2012.",
                "Participation in open-source projects and tech meetups."
            ]
        }
    }
}

base_dir = os.getcwd()

for lang, data in resumes.items():
    doc = Document()
    # Styles
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'; normal.font.size = Pt(11)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri Light'; h2.font.size = Pt(14); h2.font.bold = True
    h2.paragraph_format.space_after = Pt(6)

    # Name centered
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(data['name'])
    run.font.name = 'Calibri Light'; run.font.size = Pt(24)

    # Contact table
    table = doc.add_table(rows=1, cols=2, style='Table Grid')
    table.autofit = False
    table.columns[0].width = Inches(4)
    table.columns[1].width = Inches(2)
    cell_left, cell_right = table.rows[0].cells

    # Contacts
    for label, key, display in [
        ("Mobil", None, data['phone']),
        ("email", 'email', data['email']),
        ("LinkedIn", 'linkedin', 'LinkedIn'),
        ("GitHub", 'github', 'GitHub')
    ]:
        p = cell_left.add_paragraph()
        if key:
            p.add_run(f"{label}: ")
            url = data[key] if key != 'email' else f"mailto:{data[key]}"
            add_hyperlink(p, display, url)
        else:
            p.add_run(f"{label}: {display}")

    # Photo
    photo = os.path.join(base_dir, 'photo.jpg')
    if os.path.exists(photo):
        cell_right.add_paragraph().add_run().add_picture(photo, width=Inches(1.5))

    doc.add_paragraph()  # space after table

    # Summary
    doc.add_paragraph(data['summary_heading'], style='Heading 2')
    doc.add_paragraph(data['summary'])

    # Sections
    for section in data['sections_order']:
        doc.add_paragraph(section, style='Heading 2')
        for item in data['sections'][section]:
            para = doc.add_paragraph(item, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)

    # Save
    out = os.path.join(base_dir, f"resume_{lang}.docx")
    doc.save(out)
    print(f"{lang.upper()} resume saved: {out}")
