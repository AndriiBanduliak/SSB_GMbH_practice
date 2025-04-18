import os
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a Word paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), "0000FF"); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr)
    text_elem = OxmlElement('w:t'); text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def add_hr_word(doc):
    """Add horizontal rule in a Word document."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

# Resume content for German and English
resumes = {
    "de": {
        "filename": "resume_de",
        "name": "Andrii Banduliak",
        "contacts": [
            ("Mobil:", "+49 160 91441767", None),
            ("E‑Mail:", "aobanduliak@gmail.com", "mailto:aobanduliak@gmail.com"),
            ("LinkedIn:", "linkedin.com/in/andrii-banduliak", "https://linkedin.com/in/andrii-banduliak"),
            ("GitHub:", "github.com/AndriiBanduliak", "https://github.com/AndriiBanduliak")
        ],
        "summary_title": "Profil",
        "summary": (
            "Motivierter Softwareentwickler und Data Engineer mit juristischem Hintergrund, "
            "spezialisiert auf Python, Cloud Computing (GCP) und MLOps. Leidenschaftlich "
            "datengetriebene Lösungen und Automatisierung vorantreibend, untermauert durch "
            "umfangreiche Erfahrung in rechtlicher Recherche und Vertragsmanagement."
        ),
        "sections": [
            ("Kernkompetenzen", [
                ("Programmiersprachen:", "Python, JavaScript, SQL, Java"),
                ("Cloud & Infrastruktur:", "Google Cloud Platform (Compute Engine, BigQuery, Cloud Composer), Docker, Kubernetes"),
                ("Data Engineering & MLOps:", "Apache Airflow, Kubeflow, ML-Pipelines"),
                ("Frameworks & Tools:", "TensorFlow, PyTorch, Django, Flask, Git, GitHub Actions, Jenkins"),
                ("Datenbanken:", "PostgreSQL, MySQL, MongoDB"),
                ("Soft Skills:", "Analytisches Denken, Projektmanagement, interdisziplinäre Kommunikation")
            ]),
            ("Berufserfahrung", [
                ("Freelance Softwareentwickler & Data Engineer", "April 2023 – heute", [
                    "Entwicklung und Wartung von ML-Pipelines mit Kubeflow auf GCP für Datenvorverarbeitung und Modelltraining.",
                    "Automatisierung von Workflows mittels Apache Airflow und GitHub Actions, wodurch manuelle Aufwände um 30 % reduziert wurden.",
                    "Umsetzung eines Discord-Bots (X-Discord-bot-v-1) für Moderation und Echtzeitdatenanalyse.",
                    "Implementierung von PDF-zu-Audio-Konvertierung (pdf_to_mp3) und Speech-to-Text-Funktionalität."
                ]),
                ("Mitbegründer & Rechtsanwalt – Ledovskoy & Banduliak, Ukraine", "Mai 2018 – Mai 2023", [
                    "Leitung der Kanzlei und Beratung zu Zivil-, Handels- und Verwaltungsrecht.",
                    "Entwurf, Verhandlung und Prüfung komplexer Verträge; Sicherstellung von Compliance und Risikominimierung."
                ])
            ]),
            ("Ausgewählte Projekte", [
                "SSB_GMbH_practice (Python): Automated Analysis Platform für Geschäftsprozesse.",
                "find_word_by_lvl_in_text (Python): Semantische Textanalyse zur Wortklassifizierung.",
                "helper_to_learn_lang (Python/Django): Web-App für das Erlernen von Sprachen.",
                "mlops-on-gcp (Jupyter Notebooks): Erweiterung und Anpassung von Kubeflow-Pipelines auf GCP."
            ]),
            ("Bildung", [
                ("Master of Laws (LL.M.) – Nationale Akademie für Innere Angelegenheiten, Ukraine", "2011 – 2013")
            ]),
            ("Zertifikate", [
                "Google Cloud Digital Leader (2024)",
                "IBM Data Science Professional Certificate (2024)",
                "Meta Back-End Developer Professional Certificate (2024)",
                "Coursera: Data Engineering Foundations, GANs, Cloud Engineer Professional (2024)",
                "edX: CS50, Introduction to Computer Science with Python (2023)"
            ]),
            ("Sprachen", [
                "Ukrainisch, Russisch (Muttersprache)",
                "Englisch (fortgeschritten)",
                "Deutsch (A2)"
            ]),
            ("Schwerbehinderung", [
                "Grad der Behinderung: 80 % G B aG"
            ]),
            ("Ehrenamt & Interessen", [
                "Freiwillige Unterstützung bei Präsidentschaftswahlen (2014) und EURO 2012.",
                "Engagement in Open-Source-Projekten und Teilnahme an Tech-Meetups."
            ])
        ]
    },
    "en": {
        "filename": "resume_en",
        "name": "Andrii Banduliak",
        "contacts": [
            ("Mobile:", "+49 160 91441767", None),
            ("Email:", "aobanduliak@gmail.com", "mailto:aobanduliak@gmail.com"),
            ("LinkedIn:", "linkedin.com/in/andrii-banduliak", "https://linkedin.com/in/andrii-banduliak"),
            ("GitHub:", "github.com/AndriiBanduliak", "https://github.com/AndriiBanduliak")
        ],
        "summary_title": "Profile",
        "summary": (
            "Motivated Software Developer and Data Engineer with a legal background, "
            "specializing in Python, Google Cloud Platform (GCP), and MLOps. Passionate "
            "about data-driven solutions and automation, backed by extensive experience "
            "in legal research and contract management."
        ),
        "sections": [
            ("Core Competencies", [
                ("Programming Languages:", "Python, JavaScript, SQL, Java"),
                ("Cloud & Infrastructure:", "Google Cloud Platform (Compute Engine, BigQuery, Cloud Composer), Docker, Kubernetes"),
                ("Data Engineering & MLOps:", "Apache Airflow, Kubeflow, ML Pipelines"),
                ("Frameworks & Tools:", "TensorFlow, PyTorch, Django, Flask, Git, GitHub Actions, Jenkins"),
                ("Databases:", "PostgreSQL, MySQL, MongoDB"),
                ("Soft Skills:", "Analytical Thinking, Project Management, Cross-disciplinary Communication")
            ]),
            ("Work Experience", [
                ("Freelance Software Developer & Data Engineer", "April 2023 – Present", [
                    "Developed and maintained ML pipelines using Kubeflow on GCP for data preprocessing and model training.",
                    "Automated workflows with Apache Airflow and GitHub Actions, reducing manual tasks by 30%.",
                    "Implemented a Discord bot (X-Discord-bot-v-1) for real-time moderation and data analysis.",
                    "Integrated PDF-to-audio conversion and speech-to-text capabilities."
                ]),
                ("Co-founder & Attorney – Ledovskoy & Banduliak, Ukraine", "May 2018 – May 2023", [
                    "Led the law firm and advised on civil, commercial, and administrative law matters.",
                    "Drafted and negotiated complex contracts; ensured compliance and mitigated risks."
                ])
            ]),
            ("Selected Projects", [
                "SSB_GMbH_practice (Python): Automated business process analysis platform.",
                "find_word_by_lvl_in_text (Python): Semantic text analysis for word classification.",
                "helper_to_learn_lang (Python/Django): Web app for language learning assistance.",
                "mlops-on-gcp (Jupyter Notebooks): Adapted and extended Kubeflow pipelines on GCP."
            ]),
            ("Education", [
                ("Master of Laws (LL.M.) – National Academy of Internal Affairs, Ukraine", "2011 – 2013")
            ]),
            ("Certificates", [
                "Google Cloud Digital Leader (2024)",
                "IBM Data Science Professional Certificate (2024)",
                "Meta Back-End Developer Professional Certificate (2024)",
                "Coursera: Data Engineering Foundations, GANs, Cloud Engineer Professional (2024)",
                "edX: CS50, Introduction to Computer Science with Python (2023)"
            ]),
            ("Languages", [
                "Ukrainian, Russian (native)",
                "English (advanced)",
                "German (A2)"
            ]),
            ("Disability", [
                "Degree of disability: 80% (severe)"
            ]),
            ("Volunteer & Interests", [
                "Volunteered at presidential elections (2014) and EURO 2012.",
                "Participation in open-source projects and tech meetups."
            ])
        ]
    }
}

base_dir = os.getcwd()
styles_pdf = getSampleStyleSheet()
bullet_style = ParagraphStyle('Bullet', parent=styles_pdf['Normal'], leftIndent=12, bulletFontSize=10)

for lang, data in resumes.items():
    # Word generation
    doc = Document()
    normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(11)
    title = doc.add_paragraph(); run = title.add_run(data['name']); run.bold = True; run.font.size = Pt(16)
    add_hr_word(doc)
    # Contacts
    p = doc.add_paragraph(); run = p.add_run("Kontakte" if lang=='de' else "Contacts"); run.bold = True; run.font.size = Pt(12)
    for label, val, link in data['contacts']:
        p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent = Pt(12)
        p_run = p.add_run(label + " ")
        if link:
            add_hyperlink(p, val, link)
        else:
            p.add_run(val)
    add_hr_word(doc)
    # Summary
    p = doc.add_paragraph(); run = p.add_run(data['summary_title']); run.bold = True; run.font.size = Pt(12)
    doc.add_paragraph(data['summary'])
    add_hr_word(doc)
    # Sections
    for sec_title, content in data['sections']:
        p = doc.add_paragraph(); run = p.add_run(sec_title); run.bold = True; run.font.size = Pt(12)
        if isinstance(content, list) and all(isinstance(i, tuple) and len(i)==3 for i in content):
            # experience with title, date, bullets
            for title_txt, date_txt, bullets in content:
                p = doc.add_paragraph(); run = p.add_run(title_txt); run.bold=True; run.italic=True; run.font.size=Pt(11)
                p = doc.add_paragraph(); p.paragraph_format.left_indent=Pt(12)
                run = p.add_run(date_txt); run.italic=True; run.font.size=Pt(10)
                for b in bullets:
                    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Pt(18)
                    p.add_run(b)
        elif isinstance(content, list) and all(isinstance(i, tuple) and len(i)==2 for i in content):
            # simple label-description pairs
            for label_txt, desc in content:
                p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Pt(12)
                run = p.add_run(label_txt + " "); run.bold=True
                p.add_run(desc)
        else:
            # simple list of strings
            for item in content:
                p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Pt(12)
                p.add_run(item)
        add_hr_word(doc)
    word_path = os.path.join(base_dir, f"{data['filename']}.docx")
    doc.save(word_path)
    print(f"Saved Word: {word_path}")

    # PDF generation
    pdf_path = os.path.join(base_dir, f"{data['filename']}.pdf")
    doc_pdf = SimpleDocTemplate(pdf_path, pagesize=A4)
    story = []
    # Title
    story.append(Paragraph(data['name'], styles_pdf['Heading1']))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.black, spaceBefore=6, spaceAfter=6))
    # Contacts
    story.append(Paragraph("Kontakte" if lang=='de' else "Contacts", styles_pdf['Heading2']))
    for label, val, _ in data['contacts']:
        story.append(Paragraph(f"<b>{label}</b> {val}", styles_pdf['Normal']))
    story.append(Spacer(1, 6))
    # Summary
    story.append(Paragraph(data['summary_title'], styles_pdf['Heading2']))
    story.append(Paragraph(data['summary'], styles_pdf['Normal']))
    story.append(Spacer(1, 6))
    # Sections
    for sec_title, content in data['sections']:
        story.append(Paragraph(sec_title, styles_pdf['Heading2']))
        if isinstance(content, list) and all(isinstance(i, tuple) and len(i)==3 for i in content):
            for title_txt, date_txt, bullets in content:
                story.append(Paragraph(f"<b>{title_txt}</b>", styles_pdf['Normal']))
                story.append(Paragraph(f"<i>{date_txt}</i>", styles_pdf['Normal']))
                for b in bullets:
                    story.append(Paragraph(b, bullet_style, bulletText='•'))
        elif isinstance(content, list) and all(isinstance(i, tuple) and len(i)==2 for i in content):
            for label_txt, desc in content:
                story.append(Paragraph(f"<b>{label_txt}</b> {desc}", styles_pdf['Normal']))
        else:
            for item in content:
                story.append(Paragraph(item, bullet_style, bulletText='•'))
        story.append(Spacer(1, 6))
    doc_pdf.build(story)
    print(f"Saved PDF: {pdf_path}")
